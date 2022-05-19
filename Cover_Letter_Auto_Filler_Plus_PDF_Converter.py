import docx
from docx2pdf import convert
import os
from PyPDF2 import PdfFileMerger

#the bracket should be the file name (docx file).

doc = docx.Document('')

#remove the file if it exists

if os.path.exists(''):
    os.remove('')
    
print('Please enter the company name')
newCompany = input()

print('Please enter the street address of the company (not including city, province, and zip')
newStreetAddress = input()

print('Please enter "City, Province"')
newCityAndProvince = input()

print('Please enter the zip code')
newZip = input()

print('Please enter the name of the position you are applying for')
newPosition = input()

def docx_find_replace_text(doc, search_text, replace_text):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if search_text[search_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue

                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                    # check sequence
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            # no match so must be false positive
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        # found all chars in search_text
                        found_all = True
                        break

                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(replace_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text
            # print(p.text)


docx_find_replace_text(doc, 'COMPANYNAME', newCompany)
docx_find_replace_text(doc, 'STREETADDRESS', newStreetAddress)
docx_find_replace_text(doc, 'CITYANDPROVINCE', newCityAndProvince)
docx_find_replace_text(doc, 'ZIP', newZip)
docx_find_replace_text(doc, 'POSITION', newPosition)

#enter the file name here
doc.save('.docx')

print('When I made this program, I did not put a safeguard in case there is no address (remote company). It will likely leave a blank space')
print('Please open the word file before continuing. Enter anything to continue...')

#this section of the program is meant to convert the word files to pdf, and then combine them if you want. Enter the file names as indicated

continueProcess = input()
if os.path.exists(".pdf"):
    os.remove(".pdf")
if os.path.exists(".pdf"):
    os.remove(".pdf")
if os.path.exists(".pdf"):
    os.remove(".pdf")
convert(".docx")
convert(".docx")

print('Do you want to combine your resume and cover letter into one file? (y for yes and n for no')
combineAnswer = input()
pdfs = ["resume.pdf", "Cover_Letter.pdf"]
merger = PdfFileMerger()
if combineAnswer == 'y':
    for i in pdfs:
        merger.append(i)
    merger.write('Application.pdf')
    merger.close()
    print('Done')
else:
    print('Done')

input("prompt: ")



